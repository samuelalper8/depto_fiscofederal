"""
ConPrev — Gerador EFD-Reinf  ·  SaaS Premium (v10.3 - Master Version)
=============================================================
Código Completo: IA Gemini (Otimizada), Lógica de Salvamento (Adicionar/Sobrepor),
Ocultação de Index, Agrupamento Hierárquico no Word, PDFs e CSS Glassmorphism.
"""
import streamlit as st
import io
import os
import json
import subprocess
import tempfile
import re
from datetime import datetime, timedelta
from typing import Dict, List, Any, Optional, Tuple
from collections import defaultdict

from openpyxl import load_workbook
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Importação da IA do Google
try:
    import google.generativeai as genai
    IA_DISPONIVEL = True
except ImportError:
    IA_DISPONIVEL = False

# ── Configuração da Página ────────────────────────────────────────────────────
st.set_page_config(
    page_title="ConPrev — EFD-Reinf",
    page_icon="🌌", layout="wide",
    initial_sidebar_state="collapsed",
)

# 🔴 BLINDAGEM DE ESTADO & TEMA 🔴
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False
if "dark_mode" not in st.session_state:
    st.session_state["dark_mode"] = False
if "ia_dados_importados" not in st.session_state:
    st.session_state["ia_dados_importados"] = []

# ── Bancos de Dados Locais (JSON / NoSQL) ─────────────────────────────────────
ARQUIVO_CLIENTES = "clientes.json"
ARQUIVO_LANCAMENTOS = "lancamentos.json"

CLIENTES_PADRAO: Dict[str, Dict[str, str]] = {
    "Legislativo - Eldorado": {"UF": "MS", "CNPJ": "70.524.376/0001-80"},
    "Município - Uirapuru": {"UF": "GO", "CNPJ": "37.622.164/0001-60"},
    "Município - Santa Maria do Tocantins": {"UF": "TO", "CNPJ": "37.421.039/0001-92"},
    "Município - Lajeado": {"UF": "TO", "CNPJ": "37.420.650/0001-04"},
    "Município - Jaú do Tocantins": {"UF": "TO", "CNPJ": "37.344.413/0001-01"},
    "Município - Alcinópolis": {"UF": "MS", "CNPJ": "37.226.651/0001-04"},
    "Município - Teresina de Goiás": {"UF": "GO", "CNPJ": "25.105.339/0001-83"},
    "Município - Goianorte": {"UF": "TO", "CNPJ": "25.086.612/0001-70"},
    "Município - Palmeiras do Tocantins": {"UF": "TO", "CNPJ": "25.064.056/0001-30"},
    "Município - Maurilândia do Tocantins": {"UF": "TO", "CNPJ": "25.064.015/0001-44"},
    "Município - São Valério": {"UF": "TO", "CNPJ": "25.043.449/0001-68"},
    "Legislativo - Rio Verde": {"UF": "GO", "CNPJ": "25.040.627/0001-05"},
    "Município - Perolândia": {"UF": "GO", "CNPJ": "24.859.324/0001-48"},
    "Município - Rio Quente": {"UF": "GO", "CNPJ": "24.852.675/0001-27"},
    "Município - Sonora": {"UF": "MS", "CNPJ": "24.651.234/0001-67"},
    "Município - Chapadão do Sul": {"UF": "MS", "CNPJ": "24.651.200/0001-72"},
    "Município - Japorã": {"UF": "MS", "CNPJ": "15.905.342/0001-28"},
    "Autarquia - SAAE de Jaraguari": {"UF": "MS", "CNPJ": "15.435.936/0001-12"},
    "RPPS - Baliza": {"UF": "GO", "CNPJ": "11.329.148/0001-90"},
    "RPPS - Piranhas": {"UF": "GO", "CNPJ": "07.578.154/0001-04"},
    "RPPS - Serranópolis": {"UF": "GO", "CNPJ": "05.433.433/0001-54"},
    "RPPS - Itaberaí": {"UF": "GO", "CNPJ": "05.370.217/0001-07"},
    "Autarquia - Goiatuba IAG": {"UF": "GO", "CNPJ": "05.098.663/0001-04"},
    "RPPS - Trindade": {"UF": "GO", "CNPJ": "05.015.173/0001-05"},
    "RPPS - Barro Alto": {"UF": "GO", "CNPJ": "05.004.744/0001-06"},
    "RPPS - Crixás": {"UF": "GO", "CNPJ": "04.739.716/0001-66"},
    "Legislativo - Ceres": {"UF": "GO", "CNPJ": "04.340.201/0001-99"},
    "RPPS - Sonora": {"UF": "MS", "CNPJ": "04.318.288/0001-06"},
    "Município - Sete Quedas": {"UF": "MS", "CNPJ": "03.889.011/0001-62"},
    "Município - Tacuru": {"UF": "MS", "CNPJ": "03.888.989/0001-00"},
    "Município - Iguatemi": {"UF": "MS", "CNPJ": "03.568.318/0001-61"},
    "Município - Coxim": {"UF": "MS", "CNPJ": "03.510.211/0001-62"},
    "Município - Jaraguari": {"UF": "MS", "CNPJ": "03.501.533/0001-45"},
    "Município - Anastácio": {"UF": "MS", "CNPJ": "03.452.307/0001-11"},
    "Município - Brejinho de Nazaré": {"UF": "TO", "CNPJ": "02.884.153/0001-74"},
    "Município - Pilar de Goiás": {"UF": "GO", "CNPJ": "02.647.303/0001-26"},
    "Município - São Francisco de Goiás": {"UF": "GO", "CNPJ": "02.468.437/0001-80"},
    "Município - Itaberaí": {"UF": "GO", "CNPJ": "02.451.938/0001-53"},
    "Município - Peixe": {"UF": "TO", "CNPJ": "02.396.166/0001-02"},
    "Município - Crixás": {"UF": "GO", "CNPJ": "02.382.067/0001-63"},
    "Município - Barro Alto": {"UF": "GO", "CNPJ": "02.355.675/0001-89"},
    "Legislativo - Itapaci": {"UF": "GO", "CNPJ": "02.353.368/0001-69"},
    "Município - Córrego do Ouro": {"UF": "GO", "CNPJ": "02.321.115/0001-03"},
    "Município - São Luís de Montes Belos": {"UF": "GO", "CNPJ": "02.320.406/0001-87"},
    "Município - Goiás": {"UF": "GO", "CNPJ": "02.295.772/0001-23"},
    "Legislativo - Perolândia": {"UF": "GO", "CNPJ": "02.254.179/0001-39"},
    "Legislativo - Jaraguari": {"UF": "MS", "CNPJ": "02.210.819/0001-09"},
    "Município - Pedro Afonso": {"UF": "TO", "CNPJ": "02.070.589/0001-20"},
    "Município - Guaraí": {"UF": "TO", "CNPJ": "02.070.548/0001-33"},
    "Município - Paranaiguara": {"UF": "GO", "CNPJ": "02.056.745/0001-06"},
    "Município - Natividade": {"UF": "TO", "CNPJ": "01.809.474/0001-41"},
    "Município - Montes Claros de Goiás": {"UF": "GO", "CNPJ": "01.767.722/0001-39"},
    "Município - Brazabrantes": {"UF": "GO", "CNPJ": "01.756.741/0001-60"},
    "Município - Goiatuba": {"UF": "GO", "CNPJ": "01.753.722/0001-80"},
    "Legislativo - São Luís de Montes Belos": {"UF": "GO", "CNPJ": "01.725.501/0001-06"},
    "Município - Aguiarnópolis": {"UF": "TO", "CNPJ": "01.634.074/0001-42"},
    "Município - Novo Gama": {"UF": "GO", "CNPJ": "01.629.276/0001-04"},
    "Município - Santa Rita do Tocantins": {"UF": "TO", "CNPJ": "01.613.127/0001-49"},
    "Município - Bandeirantes do Tocantins": {"UF": "TO", "CNPJ": "01.612.819/0001-72"},
    "Município - Barra do Ouro": {"UF": "TO", "CNPJ": "01.612.818/0001-28"},
    "Autarquia - Goiatuba FESG": {"UF": "GO", "CNPJ": "01.494.665/0001-61"},
    "Município - Amaralina": {"UF": "GO", "CNPJ": "01.492.098/0001-04"},
    "Legislativo - Peixe": {"UF": "TO", "CNPJ": "01.447.812/0001-42"},
    "Município - Buriti Alegre": {"UF": "GO", "CNPJ": "01.345.909/0001-44"},
    "Município - Serranópolis": {"UF": "GO", "CNPJ": "01.343.086/0001-18"},
    "Município - Rianápolis": {"UF": "GO", "CNPJ": "01.300.094/0001-87"},
    "Município - Jaraguá": {"UF": "GO", "CNPJ": "01.223.916/0001-73"},
    "Município - Trindade": {"UF": "GO", "CNPJ": "01.217.538/0001-15"},
    "Município - Pium": {"UF": "TO", "CNPJ": "01.189.497/0001-09"},
    "Município - Piranhas": {"UF": "GO", "CNPJ": "01.168.145/0001-69"},
    "Município - Caiapônia": {"UF": "GO", "CNPJ": "01.164.946/0001-56"},
    "Município - Almas": {"UF": "TO", "CNPJ": "01.138.551/0001-89"},
    "Município - Cristalina": {"UF": "GO", "CNPJ": "01.138.122/0001-01"},
    "Município - Itapaci": {"UF": "GO", "CNPJ": "01.134.808/0001-24"},
    "Município - Ceres": {"UF": "GO", "CNPJ": "01.131.713/0001-57"},
    "Município - Corumbá de Goiás": {"UF": "GO", "CNPJ": "01.118.850/0001-51"},
    "Município - Hidrolina": {"UF": "GO", "CNPJ": "01.067.230/0001-30"},
    "Município - Cristalândia": {"UF": "TO", "CNPJ": "01.067.156/0001-52"},
    "Município - Baliza": {"UF": "GO", "CNPJ": "01.067.131/0001-59"},
    "Município - Bela Vista de Goiás": {"UF": "GO", "CNPJ": "01.005.917/0001-41"},
    "Legislativo - Costa Rica": {"UF": "MS", "CNPJ": "00.991.547/0001-04"},
    "Legislativo - Catalão": {"UF": "GO", "CNPJ": "00.833.942/0001-50"},
    "Município - Paraíso do Tocantins": {"UF": "TO", "CNPJ": "00.299.180/0001-54"},
    "Município - Campinaçu": {"UF": "GO", "CNPJ": "00.145.789/0001-79"},
    "Município - Silvanópolis": {"UF": "TO", "CNPJ": "00.114.819/0001-80"},
    "Município - Palmeirópolis": {"UF": "TO", "CNPJ": "00.007.401/0001-73"}
}

RESPONSAVEIS: Dict[str, str] = {
    "Wênia Rodrigues": "1024", "Aline Moreno": "1021",
    "Gustavo Nogueira": "1023", "Rafael Reis": "1022", "Samuel Almeida": "1020"
}

def carregar_clientes() -> dict:
    if os.path.exists(ARQUIVO_CLIENTES):
        try:
            with open(ARQUIVO_CLIENTES, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    with open(ARQUIVO_CLIENTES, "w", encoding="utf-8") as f: json.dump(CLIENTES_PADRAO, f, ensure_ascii=False, indent=4)
    return CLIENTES_PADRAO

def salvar_novo_cliente(nome: str, uf: str, cnpj: str):
    clientes = carregar_clientes()
    clientes[nome] = {"UF": uf.upper(), "CNPJ": cnpj}
    with open(ARQUIVO_CLIENTES, "w", encoding="utf-8") as f: json.dump(clientes, f, ensure_ascii=False, indent=4)

def carregar_lancamentos() -> dict:
    if os.path.exists(ARQUIVO_LANCAMENTOS):
        try:
            with open(ARQUIVO_LANCAMENTOS, "r", encoding="utf-8") as f: return json.load(f)
        except: pass
    return {}

def salvar_lancamentos(cliente: str, competencia: str, dados: list, modo: str = "sobrepor"):
    db = carregar_lancamentos()
    if cliente not in db: db[cliente] = {}
    
    if modo == "adicionar" and competencia in db[cliente]:
        db[cliente][competencia].extend(dados)
    else:
        db[cliente][competencia] = dados
        
    with open(ARQUIVO_LANCAMENTOS, "w", encoding="utf-8") as f: json.dump(db, f, ensure_ascii=False, indent=4)

# ── Engine de Temas CSS Dinâmico ──────────────────────────────────────────────
def injetar_css():
    is_dark = st.session_state["dark_mode"]
    
    if is_dark:
        bg_color = "#0B1E33"
        text_color = "#dce8f2"
        heading_color = "#fff"
        glass_bg = "rgba(255, 255, 255, 0.03)"
        glass_border = "rgba(255, 255, 255, 0.08)"
        label_color = "#7a95ad"
        card_bg = "rgba(255,255,255,0.02)"
        shadow = "rgba(0,0,0,0.3)"
    else:
        bg_color = "#F8FAFC"
        text_color = "#2D3748"
        heading_color = "#1A365D"
        glass_bg = "rgba(255, 255, 255, 0.8)"
        glass_border = "rgba(0, 0, 0, 0.08)"
        label_color = "#4A5568"
        card_bg = "#FFFFFF"
        shadow = "rgba(0,0,0,0.05)"

    css = f"""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&family=Space+Grotesk:wght@500;700&display=swap');
    .stApp {{
        background: radial-gradient(circle at 15% 50%, rgba(45, 143, 212, 0.08), transparent 25%),
                    radial-gradient(circle at 85% 30%, rgba(242, 159, 5, 0.08), transparent 25%), {bg_color} !important;
    }}
    html, body, p, span, div, label, li {{ font-family: 'Inter', sans-serif !important; color: {text_color}; }}
    h1, h2, h3, h4, h5, h6 {{ font-family: 'Space Grotesk', sans-serif !important; letter-spacing: -0.5px; color: {heading_color} !important; }}
    @keyframes fadeSlideUp {{ 0% {{ opacity: 0; transform: translateY(20px); }} 100% {{ opacity: 1; transform: translateY(0); }} }}
    .block-container {{ animation: fadeSlideUp 0.8s cubic-bezier(0.16, 1, 0.3, 1) forwards; padding-top: 2rem !important; max-width: 1100px !important; }}
    .stTextInput>div>div>input, .stDateInput>div>div>input, .stNumberInput>div>div>input, [data-baseweb="select"]>div {{
        background: {glass_bg} !important; backdrop-filter: blur(10px) !important;
        border: 1px solid {glass_border} !important; border-radius: 12px !important; color: {text_color} !important; transition: all 0.3s ease !important;
        box-shadow: 0 2px 5px rgba(0,0,0,0.02) !important;
    }}
    .stTextInput>div>div>input:focus, .stDateInput>div>div>input:focus, [data-baseweb="select"]>div:focus-within {{
        border-color: rgba(45, 143, 212, 0.5) !important; box-shadow: 0 0 15px rgba(45, 143, 212, 0.15) !important; background: {card_bg} !important;
    }}
    .stTextInput>label, .stSelectbox>label, .stDateInput>label, .stNumberInput>label {{
        color: {label_color} !important; font-size: 11px !important; font-weight: 600 !important; text-transform: uppercase; letter-spacing: 1px;
    }}
    button[data-baseweb="tab"] {{ background: transparent !important; color: {label_color} !important; font-family: 'Space Grotesk', sans-serif !important; border: none !important; outline: none !important; box-shadow: none !important; }}
    button[aria-selected="true"][data-baseweb="tab"] {{ color: #F29F05 !important; border-bottom: 2px solid #F29F05 !important; }}
    .custom-card {{
        background: {card_bg}; backdrop-filter: blur(20px); border: 1px solid {glass_border}; 
        padding: 30px; border-radius: 16px; text-align:center; box-shadow: 0 10px 30px {shadow};
    }}
    .section-card {{
        display:flex;align-items:center;gap:10px;padding:13px 18px 11px;background:{card_bg};
        border:1px solid {glass_border};border-left:3px solid #F29F05;border-radius:10px;margin-bottom:15px;box-shadow:0 2px 10px {shadow};
    }}
    .stButton>button[kind="primary"] {{
        background: linear-gradient(135deg, #F29F05, #d78904) !important; color: #FFFFFF !important; font-weight: 700 !important; font-family: 'Space Grotesk', sans-serif !important; border: none !important; border-radius: 10px !important; padding: 10px 24px !important; box-shadow: 0 4px 15px rgba(242, 159, 5, 0.3) !important; transition: all 0.3s ease !important;
    }}
    .stButton>button[kind="primary"]:hover {{ transform: translateY(-2px) scale(1.02) !important; box-shadow: 0 8px 25px rgba(242, 159, 5, 0.4) !important; }}
    .stCheckbox>label {{ color: {text_color} !important; font-size: 13px !important; cursor: pointer; }}
    #MainMenu, footer, [data-testid="stDecoration"], [data-testid="stToolbar"] {{ display: none !important; }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

injetar_css()

# ── Cérebro de IA (Gemini Vision) ─────────────────────────────────────────────
def formatar_cnpj(cnpj_str: str) -> str:
    digits = re.sub(r'\D', '', str(cnpj_str))
    if len(digits) == 14:
        return f"{digits[:2]}.{digits[2:5]}.{digits[5:8]}/{digits[8:12]}-{digits[12:]}"
    return cnpj_str

def extrair_dados_ia_gemini(uploaded_file, api_key: str) -> Optional[Dict[str, Any]]:
    if not IA_DISPONIVEL:
        st.error("Biblioteca 'google-generativeai' não instalada no servidor.")
        return None
        
    genai.configure(api_key=api_key)
    
    prompt = """
    Analise esta imagem ou PDF de Nota Fiscal de Serviço (NFS-e).
    Extraia os seguintes dados tributários e me devolva ESTRITAMENTE um objeto JSON válido, sem NENHUM texto adicional ou markdown, com as chaves exatas abaixo:
    {
        "Órgão": "Nome do TOMADOR do serviço",
        "CNPJ Tomador": "Apenas os números do CNPJ do tomador",
        "Nº NF": "Número da Nota",
        "CNPJ Prestador": "Apenas os números do CNPJ do prestador",
        "Total Contrib. Prev.": 0.0
    }
    Se não houver INSS retido, retorne 0.0 no Total Contrib. Prev.
    """
    
    try:
        ext = os.path.splitext(uploaded_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
            
        sample_file = genai.upload_file(path=tmp_path)
        
        # Uso do modelo mais otimizado e resiliente para a cota gratuita
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content([prompt, sample_file])
        
        genai.delete_file(sample_file.name)
        os.remove(tmp_path)
        
        txt_limpo = response.text.replace('```json', '').replace('```', '').strip()
        dados = json.loads(txt_limpo)
        
        dados["CNPJ Tomador"] = formatar_cnpj(dados.get("CNPJ Tomador", ""))
        dados["CNPJ Prestador"] = formatar_cnpj(dados.get("CNPJ Prestador", ""))
        dados["Total Contrib. Prev."] = float(str(dados.get("Total Contrib. Prev.", 0.0)).replace(',', '.'))
        dados["Compensação"] = 0.0
        
        return dados
    except json.JSONDecodeError:
        st.error("A IA não conseguiu estruturar os dados. A imagem pode estar ilegível.")
        return None
    except Exception as e:
        if "429" in str(e):
            st.error("⚠️ Limite de requisições gratuitas do Google excedido. Aguarde cerca de 1 minuto para tentar novamente.")
        else:
            st.error(f"Erro de processamento da IA: {e}")
        return None

# ── Funções Auxiliares ────────────────────────────────────────────────────────
def get_datas_padrao() -> Tuple[str, str, str, datetime]:
    hoje = datetime.now()
    primeiro_dia = hoje.replace(day=1)
    ultimo_dia_mes_ant = primeiro_dia - timedelta(days=1)
    meses_pt = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
    comp_folha = f"{meses_pt[ultimo_dia_mes_ant.month - 1]}/{ultimo_dia_mes_ant.year}"
    comp_email = f"{ultimo_dia_mes_ant.month:02d}/{ultimo_dia_mes_ant.year}"
    venc_dt = datetime(hoje.year, hoje.month, 20)
    venc_str = venc_dt.strftime("%d/%m/%Y")
    return comp_folha, venc_str, comp_email, venc_dt

def safe_float(value: Any) -> float:
    if value is None: return 0.0
    try: return float(str(value).replace(',', '.')) if isinstance(value, str) else float(value)
    except (ValueError, TypeError): return 0.0

def _brl_fmt(valor: Any) -> str:
    return f"R$ {safe_float(valor):,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')

# ── Funções de Manipulação Word & PDF ─────────────────────────────────────────
def set_cell_background(cell, fill_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def criar_tabela_reinf(doc: Document, dados_nfs: List[Dict[str, Any]]) -> Any:
    headers = ['Órgão', 'CNPJ Tomador', 'Nº NF', 'CNPJ Prestador', 'Total Contrib. Prev.', 'Compensação']
    
    if not dados_nfs:
        table = doc.add_table(rows=2, cols=6)
        table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]; set_cell_background(cell, "D9D9D9")
            p = cell.paragraphs[0]; r = p.add_run(h); r.font.bold = True; r.font.size = Pt(10); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_msg = table.rows[1].cells
        row_msg[0].text = "Nenhuma retenção de INSS declarada na EFD-REINF"
        row_msg[0].merge(row_msg[5]); row_msg[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return table
        
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; set_cell_background(cell, "D9D9D9")
        p = cell.paragraphs[0]; r = p.add_run(h); r.font.bold = True; r.font.size = Pt(10); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    grupos = defaultdict(lambda: defaultdict(list))
    for nf in dados_nfs:
        orgao = str(nf.get('Órgão', 'Não Informado')).strip()
        if not orgao or orgao.lower() == 'none': orgao = 'Não Informado'
        prestador = str(nf.get('CNPJ Prestador', 'Não Informado')).strip()
        if not prestador or prestador.lower() == 'none': prestador = 'Não Informado'
        grupos[orgao][prestador].append(nf)

    total_geral_contrib = 0.0
    total_geral_comp = 0.0

    for orgao, prestadores in grupos.items():
        subtotal_orgao_contrib = 0.0
        subtotal_orgao_comp = 0.0
        
        for prestador, nfs in prestadores.items():
            subtotal_prest_contrib = 0.0
            subtotal_prest_comp = 0.0
            
            for nf in nfs:
                row = table.add_row().cells
                row[0].text = str(nf.get('Órgão', ''))
                row[1].text = str(nf.get('CNPJ Tomador', ''))
                row[2].text = str(nf.get('Nº NF', ''))
                row[3].text = str(nf.get('CNPJ Prestador', ''))
                row[4].text = _brl_fmt(nf.get('Total Contrib. Prev.'))
                row[5].text = _brl_fmt(nf.get('Compensação'))
                
                subtotal_prest_contrib += safe_float(nf.get('Total Contrib. Prev.'))
                subtotal_prest_comp += safe_float(nf.get('Compensação'))
                
                for cell in row:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs: run.font.size = Pt(10)
            
            st_prest_row = table.add_row().cells
            for cell in st_prest_row: set_cell_background(cell, "FDFDFD")
            st_prest_row[0].text = f"Subtotal - CNPJ {prestador}"
            st_prest_row[0].merge(st_prest_row[3])
            st_prest_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            st_prest_row[4].text = _brl_fmt(subtotal_prest_contrib)
            st_prest_row[5].text = _brl_fmt(subtotal_prest_comp)
            
            for idx in [0, 4, 5]:
                cell = st_prest_row[idx]
                p = cell.paragraphs[0]
                if not p.runs: p.add_run(cell.text)
                for run in p.runs: run.bold = True; run.font.size = Pt(10)
                if idx != 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            subtotal_orgao_contrib += subtotal_prest_contrib
            subtotal_orgao_comp += subtotal_prest_comp

        st_orgao_row = table.add_row().cells
        for cell in st_orgao_row: set_cell_background(cell, "F2F2F2")
        st_orgao_row[0].text = f"Subtotal do Órgão ({orgao})"
        st_orgao_row[0].merge(st_orgao_row[3])
        st_orgao_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        st_orgao_row[4].text = _brl_fmt(subtotal_orgao_contrib)
        st_orgao_row[5].text = _brl_fmt(subtotal_orgao_comp)
        
        for idx in [0, 4, 5]:
            cell = st_orgao_row[idx]
            p = cell.paragraphs[0]
            if not p.runs: p.add_run(cell.text)
            for run in p.runs: run.bold = True; run.font.size = Pt(10)
            if idx != 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        total_geral_contrib += subtotal_orgao_contrib
        total_geral_comp += subtotal_orgao_comp

    t_row = table.add_row().cells
    for cell in t_row: set_cell_background(cell, "EAEAEA")
    t_row[0].text = "TOTAL GERAL"
    t_row[0].merge(t_row[3])
    t_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    t_row[4].text = _brl_fmt(total_geral_contrib)
    t_row[5].text = _brl_fmt(total_geral_comp)
    
    for idx in [0, 4, 5]:
        cell = t_row[idx]
        p = cell.paragraphs[0]
        if not p.runs: p.add_run(cell.text)
        for run in p.runs: run.bold = True; run.font.size = Pt(10)
        if idx != 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    return table

def replace_everywhere(doc: Document, old: str, new: str) -> None:
    def repl(par):
        if old in par.text:
            for run in par.runs:
                if old in run.text: run.text = run.text.replace(old, new)
            if old in par.text: par.text = par.text.replace(old, new)
    for p in doc.paragraphs: repl(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs: repl(p)

def converter_para_pdf(docx_bytes: bytes) -> Optional[bytes]:
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        with open(docx_path, "wb") as f: f.write(docx_bytes)
        try:
            subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, docx_path], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            pdf_path = os.path.join(tmpdir, "temp.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f: return f.read()
        except Exception: return None
    return None

# ── UI Components Principais ──────────────────────────────────────────────────
def render_login() -> None:
    _, col, _ = st.columns([1.4, 1, 1.4])
    with col:
        st.markdown("""<style> [data-testid="stSidebar"] {display: none;} </style>""", unsafe_allow_html=True)
        st.markdown("""
        <div class="custom-card" style="margin-top: 15vh;">
          <div style="width:70px;height:70px;background:linear-gradient(135deg,#F29F05,#d78904);border-radius:20px;display:inline-flex;align-items:center;justify-content:center;font-size:32px;box-shadow:0 10px 30px rgba(242,159,5,.4);margin-bottom:20px">🌌</div>
          <h2 style="font-size:28px;font-weight:800;margin:0 0 8px; font-family:'Space Grotesk', sans-serif;">ConPrev</h2>
          <p style="font-size:12px;letter-spacing:2px;text-transform:uppercase;margin:0 0 30px 0; opacity: 0.7;">EFD-Reinf &middot; Sistema de Retenções</p>
        </div>""", unsafe_allow_html=True)
        
        pwd = st.text_input("Credencial de Acesso", type="password", placeholder="••••••••", label_visibility="collapsed")
        
        if st.button("Acessar Plataforma", type="primary", use_container_width=True):
            senha_oficial = st.secrets.get("APP_PASSWORD", None)
            if not senha_oficial:
                st.error("⚠️ Infraestrutura: A variável 'APP_PASSWORD' não foi configurada nos Secrets.")
                return

            if pwd == senha_oficial:
                st.session_state["authenticated"] = True
                st.rerun()
            else:
                st.error("⚠️ Senha incorreta. Acesso negado.")

def render_header():
    left, mid, right = st.columns([6, 2, 1])
    with left:
        st.markdown("""
        <div style="display:flex;align-items:center;gap:14px;padding:6px 0 16px">
          <div style="width:42px;height:42px;flex-shrink:0;background:linear-gradient(145deg,#F29F05,#d78904);border-radius:10px;display:flex;align-items:center;justify-content:center;font-size:20px;box-shadow:0 4px 14px rgba(242,159,5,.4)">📄</div>
          <div>
            <div style="font-size:22px;font-weight:800;line-height:1.2;font-family:'Space Grotesk', sans-serif;">Folha de Rosto <span style="font-weight:400;opacity:0.7;font-size:18px;margin-left:6px">EFD-Reinf</span></div>
            <div style="font-size:11px;opacity:0.6;margin-top:2px">Automação de Documentos &nbsp;·&nbsp; Fisco Federal</div>
          </div>
        </div>""", unsafe_allow_html=True)
    with mid:
        st.markdown("<br>", unsafe_allow_html=True)
        st.session_state["dark_mode"] = st.toggle("🌙 Modo Escuro", value=st.session_state["dark_mode"])
    with right:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("↩ Sair", key="logout_btn", use_container_width=True):
            st.session_state["authenticated"] = False
            st.rerun()

def render_app():
    render_header()
    comp_folha, venc_str_padrao, comp_email, venc_dt_padrao = get_datas_padrao()
    tab1, tab2, tab3 = st.tabs(["📝 1. Lançador de Notas (Nuvem/IA)", "⚙️ 2. Gerador Oficial (Word/PDF)", "🏢 3. Gestão de Clientes"])
    
    clientes_bd = carregar_clientes()
    lancamentos_bd = carregar_lancamentos()
    
    # --- TAB 1: LANÇADOR (COM IA VISION) ---
    with tab1:
        st.markdown("<br>", unsafe_allow_html=True)
        
        with st.expander("🧠 Leitor de IA (PDFs e Imagens de Notas Fiscais)", expanded=False):
            st.markdown("<p style='font-size:12px; opacity:0.8;'>Arraste PDFs, JPGs ou PNGs de notas escaneadas/amassadas. A Inteligência Artificial Gemini irá ler a imagem e extrair os dados tributários para você.</p>", unsafe_allow_html=True)
            
            chave_gemini = st.secrets.get("GEMINI_API_KEY", None)
            
            if not chave_gemini:
                st.warning("⚠️ Chave da API do Gemini (GEMINI_API_KEY) não configurada nos Secrets. Configure para liberar o leitor.")
            else:
                arquivos_ia = st.file_uploader("Arraste fotos ou PDFs aqui", type=["pdf", "png", "jpg", "jpeg"], accept_multiple_files=True)
                
                if st.button("✨ Processar Documentos com Inteligência Artificial") and arquivos_ia:
                    novos_dados = []
                    sucessos = 0
                    
                    barra_progresso = st.progress(0)
                    for i, arq in enumerate(arquivos_ia):
                        st.toast(f"Analisando: {arq.name}...", icon='👁️')
                        dado_extraido = extrair_dados_ia_gemini(arq, chave_gemini)
                        
                        if dado_extraido:
                            novos_dados.append(dado_extraido)
                            sucessos += 1
                        barra_progresso.progress((i + 1) / len(arquivos_ia))
                        
                    if novos_dados:
                        st.session_state["ia_dados_importados"] = novos_dados
                        st.success(f"✅ {sucessos} notas lidas pela IA! Por favor, revise os dados na tabela abaixo.")

        st.markdown("<h4 style='font-size:16px; margin-top:20px; margin-bottom:20px;'>Tabela de Conferência e Salvamento na Nuvem</h4>", unsafe_allow_html=True)
        
        c1, c2 = st.columns(2)
        with c1: cliente_t1 = st.selectbox("Selecione o Cliente (Para Salvar na Nuvem)", list(clientes_bd.keys()), key="cli_t1")
        with c2: comp_t1 = st.text_input("Competência", value=comp_folha, key="comp_t1")
        
        dados_atuais = lancamentos_bd.get(cliente_t1, {}).get(comp_t1, [])
        cols = ["Órgão", "CNPJ Tomador", "Nº NF", "CNPJ Prestador", "Total Contrib. Prev.", "Compensação"]
        
        dados_tabela = dados_atuais + st.session_state.get("ia_dados_importados", [])
        
        if dados_tabela:
            df_base = pd.DataFrame(dados_tabela)
            for c in cols:
                if c not in df_base.columns: df_base[c] = None
            df_base = df_base[cols]
        else:
            df_base = pd.DataFrame(columns=cols)
            for _ in range(5): df_base.loc[len(df_base)] = [None]*6
            
        # Oculta a numeração da linha (Index)
        df_editado = st.data_editor(df_base, num_rows="dynamic", use_container_width=True, hide_index=True)
        
        colA, colB = st.columns(2)
        with colA:
            modo_salvar = st.radio("Opção de Salvamento:", ["Substituir todos os dados existentes", "Adicionar as notas aos dados existentes"], horizontal=True)
            if st.button("💾 Salvar Lançamentos na Nuvem", type="primary", use_container_width=True):
                df_limpo = df_editado.dropna(how="all").where(pd.notnull(df_editado), None)
                dados_salvar = df_limpo.to_dict(orient="records")
                
                modo = "sobrepor" if "Substituir" in modo_salvar else "adicionar"
                salvar_lancamentos(cliente_t1, comp_t1, dados_salvar, modo)
                
                st.session_state["ia_dados_importados"] = [] 
                st.toast(f"Lançamentos salvos para {cliente_t1}!", icon='☁️')
                st.rerun()
                
        with colB:
            st.markdown("<br>", unsafe_allow_html=True)
            df_export = df_editado.dropna(how="all")
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer: df_export.to_excel(writer, sheet_name="Valores", index=False)
            st.download_button("📥 Baixar Planilha Manual (.xlsx)", data=output.getvalue(), file_name="Lançamentos.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)

    with tab2:
        st.markdown("<br>", unsafe_allow_html=True)
        colL, colR = st.columns([1, 1], gap="large")
        
        with colL:
            st.markdown(f"""<div class="section-card" style="border-left-color: #F29F05"><span style="font-size:15px">⚙️</span><span style="font-size:11.5px;font-weight:700;text-transform:uppercase;letter-spacing:1.2px">Configurações do Ato</span></div>""", unsafe_allow_html=True)
            cliente_sel = st.selectbox("Selecione o Cliente", list(clientes_bd.keys()), key="cli_t2")
            
            c_ato1, c_ato2 = st.columns([2, 1])
            with c_ato1: num_ato_int = st.number_input("Nº Inicial do Ato", min_value=1, value=1, step=1)
            with c_ato2: ano_ato = st.text_input("Ano", value=str(datetime.now().year))
            num_ato = f"{num_ato_int:03d}/{ano_ato}"
            
            resp_sel = st.selectbox("Responsável", list(RESPONSAVEIS.keys()))
            competencia = st.text_input("Competência", value=comp_folha, key="comp_t2")
            vencimento = st.text_input("Vencimento", value="20/03/2026")
            tipo_darf = st.radio("Tipo de Documento", ["Reinf", "Avulso", "Sem DARF"], horizontal=True)
            
        with colR:
            st.markdown(f"""<div class="section-card" style="border-left-color: #2a9c6b"><span style="font-size:15px">📤</span><span style="font-size:11.5px;font-weight:700;text-transform:uppercase;letter-spacing:1.2px">Base de Dados</span></div>""", unsafe_allow_html=True)
            houve_retencao = st.checkbox("✅ Houve retenções a declarar?", value=True)
            
            arq_excel = None
            dados_nfs = []
            can_run = True
            
            if houve_retencao:
                fonte_dados = st.radio("Fonte dos Dados:", ["☁️ Nuvem (Lançamentos Salvos da IA/Manual)", "📂 Upload Antigo (.xlsx)"], horizontal=True)
                
                if "Nuvem" in fonte_dados:
                    dados_nfs = lancamentos_bd.get(cliente_sel, {}).get(competencia, [])
                    if dados_nfs:
                        st.success(f"✅ {len(dados_nfs)} notas carregadas automaticamente do servidor.")
                    else:
                        st.warning("⚠️ Nenhum lançamento encontrado na nuvem. Vá na aba 'Lançador' para usar a IA ou mude para Upload.")
                        can_run = False
                else:
                    arq_excel = st.file_uploader("Upload da Planilha Excel (.xlsx)", type=["xlsx"])
                    can_run = bool(arq_excel)
                    if can_run:
                        wb = load_workbook(io.BytesIO(arq_excel.getvalue()), data_only=True)
                        aba = [n for n in wb.sheetnames if n.lower().startswith("valores")][0]
                        ws = wb[aba]
                        headers = [str(c.value) for c in ws[1]]
                        dados_nfs = [dict(zip(headers, r)) for r in ws.iter_rows(min_row=2, values_only=True) if not all(c is None for c in r)]
            else:
                st.info("ℹ️ Declaração sem movimento. A grelha de lançamentos será omitida.")

            if st.button("Gerar Documentos Finais", type="primary", use_container_width=True, disabled=not can_run):
                with st.spinner("Compilando Documento com Subtotais Analíticos..."):
                    chk_reinf = "☒" if tipo_darf == "Reinf" else "☐"
                    chk_avulso = "☒" if tipo_darf == "Avulso" else "☐"
                    uf = clientes_bd[cliente_sel].get("UF", "")
                    
                    contexto = {
                        '{{numero_ato}}': num_ato, '{{data_emissao}}': datetime.now().strftime('%d/%m/%Y'),
                        '{{municipio_uf}}': f"{cliente_sel} / {uf}" if uf else cliente_sel,
                        '{{competencia}}': competencia, '{{vencimento}}': vencimento,
                        '{{responsavel}}': resp_sel, '{{ramal}}': RESPONSAVEIS[resp_sel],
                        '{{check_reinf}}': chk_reinf, '{{check_avulso}}': chk_avulso
                    }
                    
                    try:
                        with open("Modelo_Folha_Rosto.docx", "rb") as f: doc = Document(io.BytesIO(f.read()))
                        for k, v in contexto.items(): replace_everywhere(doc, k, v)
                        
                        tabela = criar_tabela_reinf(doc, dados_nfs)
                        target = next((p for p in doc.paragraphs if "{{TABELA_NOTAS}}" in p.text), None)
                        if target:
                            target._p.addnext(tabela._tbl)
                            target.text = ""
                        
                        buf_docx = io.BytesIO()
                        doc.save(buf_docx)
                        bytes_docx = buf_docx.getvalue()
                        
                        st.toast('Documento Word gerado com sucesso!', icon='🎉')
                        
                        dl1, dl2 = st.columns(2)
                        with dl1: st.download_button("📥 Baixar WORD (.docx)", data=bytes_docx, file_name=f"Folha Rosto - {cliente_sel}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                        
                        with st.spinner("Gerando PDF..."):
                            bytes_pdf = converter_para_pdf(bytes_docx)
                            with dl2:
                                if bytes_pdf: st.download_button("📥 Baixar PDF (.pdf)", data=bytes_pdf, file_name=f"Folha Rosto - {cliente_sel}.pdf", mime="application/pdf", use_container_width=True)
                                else: st.button("🚫 PDF Indisponível (Instale LibreOffice)", disabled=True, use_container_width=True)
                    except Exception as e:
                        st.error(f"❌ Erro crítico: {e}")

    with tab3:
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("<div class='custom-card' style='margin-top:0; padding: 25px; text-align: left;'>", unsafe_allow_html=True)
        st.markdown("<h4 style='font-size:18px; margin-bottom:15px;'>🏢 Cadastrar Novo Cliente no Sistema</h4>", unsafe_allow_html=True)
        
        with st.form("form_novo_cliente", clear_on_submit=True):
            cc1, cc2, cc3 = st.columns([2, 1, 1])
            with cc1: novo_nome = st.text_input("Nome (Ex: Município - São Paulo)")
            with cc2: nova_uf = st.text_input("UF (Ex: SP)", max_chars=2)
            with cc3: novo_cnpj = st.text_input("CNPJ (Com pontuação)")
            
            if st.form_submit_button("Salvar Cliente na Base de Dados", type="primary", use_container_width=True):
                if novo_nome and nova_uf and novo_cnpj:
                    salvar_novo_cliente(novo_nome, nova_uf, novo_cnpj)
                    st.toast(f'Cliente "{novo_nome}" salvo com sucesso!', icon='💾')
                    st.rerun()
                else:
                    st.error("Preencha todos os campos para cadastrar um novo cliente.")
        st.markdown("</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    if not st.session_state["authenticated"]:
        render_login()
    else:
        render_app()
